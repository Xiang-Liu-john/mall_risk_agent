from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


PROJECT_DIR = Path(__file__).resolve().parents[1]
RAG_SOURCES = PROJECT_DIR / "rag_sources"


DOCX_TARGETS = {
    "contracts/merchant_lease_contract_sample.md": "contracts/商户租赁合同样本.docx",
    "contracts/deposit_and_termination_clauses.md": "contracts/保证金扣划与解约条款.docx",
    "policies/arrears_collection_workflow.md": "policies/欠费催收流程.docx",
    "policies/risk_grading_playbook.md": "policies/风险分级处置规则.docx",
    "inspection/monthly_store_inspection_record.md": "inspection/门店巡检记录.docx",
    "inspection/customer_complaint_followup.md": "inspection/投诉处理记录.docx",
}

PDF_TARGETS = {
    "contracts/arrears_collection_terms.md": "contracts/欠费催缴条款.pdf",
    "policies/merchant_management_policy.md": "policies/商户管理制度.pdf",
    "policies/legal_approval_policy.md": "policies/法务审批制度.pdf",
    "inspection/closure_risk_checklist.md": "inspection/停业撤店检查记录.pdf",
}

TABLE_TARGETS = [
    "tables/historical_sales_sample.csv",
    "tables/arrears_ledger_sample.csv",
    "tables/deposit_ledger_sample.csv",
    "tables/lease_expiry_sample.csv",
    "tables/floor_category_benchmark.csv",
]


def clean_markdown_line(line: str) -> str:
    return line.replace("**", "").replace("`", "").strip()


def add_docx_paragraph(document: Document, raw_line: str) -> None:
    line = raw_line.strip()
    if not line:
        return

    if line.startswith("# "):
        paragraph = document.add_paragraph(clean_markdown_line(line[2:]))
        paragraph.style = document.styles["Title"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    if line.startswith("## "):
        document.add_heading(clean_markdown_line(line[3:]), level=1)
        return

    if line.startswith("### "):
        document.add_heading(clean_markdown_line(line[4:]), level=2)
        return

    if line.startswith("- "):
        document.add_paragraph(clean_markdown_line(line[2:]), style="List Bullet")
        return

    if line[0].isdigit() and ". " in line[:4]:
        document.add_paragraph(clean_markdown_line(line), style="List Number")
        return

    document.add_paragraph(clean_markdown_line(line))


def create_docx(source: Path, target: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Microsoft YaHei"
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.name = "Microsoft YaHei"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.color.rgb = RGBColor(40, 55, 75)
    styles["Heading 2"].font.name = "Microsoft YaHei"
    styles["Heading 2"].font.size = Pt(12)

    for line in source.read_text(encoding="utf-8").splitlines():
        add_docx_paragraph(document, line)

    document.core_properties.author = "Mall Risk Agent"
    document.core_properties.title = source.stem.replace("_", " ")
    document.save(target)


def create_pdf(source: Path, target: Path) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    doc = SimpleDocTemplate(
        str(target),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "ChineseBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=10.5,
        leading=16,
        textColor=colors.HexColor("#1f2937"),
        spaceAfter=7,
    )
    title = ParagraphStyle(
        "ChineseTitle",
        parent=body,
        fontSize=18,
        leading=24,
        alignment=1,
        textColor=colors.HexColor("#111827"),
        spaceAfter=12,
    )
    h2 = ParagraphStyle(
        "ChineseH2",
        parent=body,
        fontSize=13,
        leading=18,
        textColor=colors.HexColor("#334155"),
        spaceBefore=8,
        spaceAfter=6,
    )

    story = []
    for raw_line in source.read_text(encoding="utf-8").splitlines():
        line = clean_markdown_line(raw_line)
        if not line:
            story.append(Spacer(1, 4))
        elif line.startswith("# "):
            story.append(Paragraph(line[2:], title))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], h2))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:], h2))
        elif line.startswith("- "):
            story.append(Paragraph(f"• {line[2:]}", body))
        else:
            story.append(Paragraph(line, body))
    doc.build(story)


def create_xlsx(source: Path, target: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "原始台账"

    with source.open("r", encoding="utf-8-sig", newline="") as file:
        rows = list(csv.reader(file))

    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row, start=1):
            cell = sheet.cell(row=row_index, column=col_index, value=value)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            if row_index == 1:
                cell.font = Font(name="Microsoft YaHei", bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="4F6F8F")
            else:
                cell.font = Font(name="Microsoft YaHei", size=10)

    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for col_index in range(1, sheet.max_column + 1):
        values = [str(sheet.cell(row=row, column=col_index).value or "") for row in range(1, sheet.max_row + 1)]
        width = min(max(max(len(value) for value in values) + 2, 12), 28)
        sheet.column_dimensions[get_column_letter(col_index)].width = width

    workbook.save(target)


def main() -> None:
    created: list[Path] = []
    for src, dst in DOCX_TARGETS.items():
        target = RAG_SOURCES / dst
        create_docx(RAG_SOURCES / src, target)
        created.append(target)

    for src, dst in PDF_TARGETS.items():
        target = RAG_SOURCES / dst
        create_pdf(RAG_SOURCES / src, target)
        created.append(target)

    for src in TABLE_TARGETS:
        source = RAG_SOURCES / src
        chinese_names = {
            "historical_sales_sample.csv": "历史销售数据样本.xlsx",
            "arrears_ledger_sample.csv": "欠费台账样本.xlsx",
            "deposit_ledger_sample.csv": "保证金台账样本.xlsx",
            "lease_expiry_sample.csv": "合同到期表样本.xlsx",
            "floor_category_benchmark.csv": "楼层业态基准表样本.xlsx",
        }
        target = source.with_name(chinese_names[source.name])
        create_xlsx(source, target)
        created.append(target)

    for path in created:
        print(path.relative_to(PROJECT_DIR))


if __name__ == "__main__":
    main()
